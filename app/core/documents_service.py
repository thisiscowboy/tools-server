import os
import io
import hashlib
import json
import tempfile
import time
import threading
import random
import uuid
import re
import logging
from pathlib import Path
from typing import List, Dict, Any, Optional, Union, Set, Tuple
from fastapi import HTTPException
from git import Repo
from pydantic import BaseModel, Field
from app.models.documents import (
    DocumentType,
    CreateDocumentRequest,
    UpdateDocumentRequest,
    DocumentResponse,
    DocumentVersionResponse,
    DocumentContentResponse,
)
from app.core.filesystem_service import FilesystemService
from app.core.memory_service import MemoryService
from app.core.git_service import GitService
from app.utils.config import get_config

logger = logging.getLogger(__name__)

try:
    import markdown
except ImportError:
    markdown = None
    logger.warning("Markdown package not installed. Markdown to HTML conversion will be limited.")


class DocumentsService:
    def __init__(self, base_path: str = None, large_content_threshold: int = 100000):
        """Initialize the document service"""
        config = get_config()

        self.base_path = Path(base_path or os.path.join(os.getcwd(), "data", "documents"))
        self.base_path.mkdir(parents=True, exist_ok=True)

        self.repo_path = str(self.base_path)
        self.git_service = GitService()

        for doc_type in DocumentType:
            (self.base_path / doc_type.value).mkdir(exist_ok=True)

        readme_path = self.base_path / "README.md"
        if not readme_path.exists():
            with open(readme_path, "w", encoding="utf-8") as f:
                f.write("# Document Storage\n\n")
                f.write(
                    "This directory contains documents managed by the Tools Server Document Service.\n"
                )
                f.write(
                    "It supports various document types including manuscripts (stories, novels),\n"
                )
                f.write("documentation (technical, research), and datasets.")
                f.write(
                    "\n\nLarge documents (70,000+ words) are automatically chunked for better performance.\n"
                )

            self.git_service.add_files(self.repo_path, ["README.md"])
            self.git_service.commit_changes(self.repo_path, "Initialize document repository")

        self.large_content_threshold = large_content_threshold

        self.index_dir = self.base_path / ".index"
        self.index_dir.mkdir(exist_ok=True)

        self.vector_search_enabled = False
        self.vector_model = None

        try:
            import numpy as np
            from sentence_transformers import SentenceTransformer

            self.np = np
            self.vector_model = SentenceTransformer("all-MiniLM-L6-v2")
            self.vector_search_enabled = True
            self.vector_index_path = self.base_path / ".vectors"
            self.vector_index_path.mkdir(exist_ok=True)
        except ImportError:
            logger.warning("Vector search dependencies not installed. Semantic search disabled.")

        self.index_lock = threading.Lock()
        self.file_locks = {}

    def _get_file_lock(self, doc_id):
        """Get a lock for a specific file to prevent concurrent modifications"""
        if doc_id not in self.file_locks:
            self.file_locks[doc_id] = threading.Lock()
        return self.file_locks[doc_id]

    def create_document(
        self,
        title: str,
        content: str,
        document_type: DocumentType,
        metadata: Dict[str, Any] = None,
        tags: List[str] = None,
        source_url: Optional[str] = None,
        storage_type: str = "local",
    ) -> Dict[str, Any]:
        """Create a new document with Git versioning"""
        doc_id = f"doc_{int(time.time())}_{uuid.uuid4().hex[:8]}"
        now = int(time.time())

        metadata = metadata or {}
        tags = tags or []

        doc_dir = self.base_path / document_type.value
        doc_path = doc_dir / f"{doc_id}.md"

        frontmatter = f"---\ntitle: {title}\ncreated_at: {now}\nupdated_at: {now}\nid: {doc_id}\ndocument_type: {document_type.value}\n"

        if tags:
            frontmatter += f"tags: {', '.join(tags)}\n"

        if source_url:
            frontmatter += f"source_url: {source_url}\n"

        for key, value in metadata.items():
            if isinstance(value, (str, int, float, bool)):
                frontmatter += f"{key}: {value}\n"

        frontmatter += "---\n\n"
        full_content = frontmatter + content

        doc_path.write_text(full_content, encoding="utf-8")

        self._update_index(
            doc_id,
            {
                "id": doc_id,
                "title": title,
                "document_type": document_type.value,
                "created_at": now,
                "updated_at": now,
                "tags": tags,
                "metadata": metadata,
                "size_bytes": len(full_content.encode("utf-8")),
                "source_url": source_url,
                "path": str(doc_path.relative_to(self.base_path)),
            },
        )

        rel_path = doc_path.relative_to(self.base_path)
        self.git_service.add_files(self.repo_path, [str(rel_path)])
        self.git_service.commit_changes(self.repo_path, f"Created document: {title}")

        self._update_memory_graph(doc_id, title, document_type.value, tags, metadata, source_url)

        self.generate_embeddings(doc_id, content)

        return self.get_document(doc_id)

    def update_document(
        self,
        doc_id: str,
        title: Optional[str] = None,
        content: Optional[str] = None,
        metadata: Optional[Dict[str, Any]] = None,
        tags: Optional[List[str]] = None,
        commit_message: str = "Updated document",
        expected_version: Optional[str] = None,
    ) -> Dict[str, Any]:
        """Update an existing document with version control"""
        with self._get_file_lock(doc_id):
            doc_info = self._get_document_index(doc_id)
            if not doc_info:
                raise ValueError(f"Document with ID {doc_id} not found")

            if expected_version:
                current_version = None
                try:
                    log = self.git_service.get_log(
                        self.repo_path, max_count=1, file_path=doc_info["path"]
                    )
                    if log.get("commits"):
                        current_version = log["commits"][0]["hash"]
                    if current_version and current_version != expected_version:
                        raise ValueError(
                            "Document has been modified since you loaded it. Please refresh and try again."
                        )
                except Exception as e:
                    logger.warning(f"Version check failed: {e}")

            doc_path = self.base_path / doc_info["path"]
            if not doc_path.exists():
                raise ValueError(f"Document file not found at {doc_path}")

            current_content = doc_path.read_text(encoding="utf-8")

            frontmatter_match = re.match(r"---(.*?)---\n\n", current_content, re.DOTALL)
            if not frontmatter_match:
                raise ValueError("Invalid document format: missing frontmatter")

            frontmatter = frontmatter_match.group(1)
            existing_content = current_content[frontmatter_match.end() :]

            frontmatter_dict = {}
            for line in frontmatter.strip().split("\n"):
                if ": " in line:
                    key, value = line.split(": ", 1)
                    frontmatter_dict[key] = value

            now = int(time.time())
            frontmatter_dict["updated_at"] = str(now)

            if title:
                frontmatter_dict["title"] = title

            if tags is not None:
                if tags:
                    frontmatter_dict["tags"] = ", ".join(tags)
                else:
                    if "tags" in frontmatter_dict:
                        del frontmatter_dict["tags"]

            if metadata:
                for key, value in metadata.items():
                    if isinstance(value, (str, int, float, bool)):
                        frontmatter_dict[key] = str(value)

            new_frontmatter = "---\n"
            for key, value in frontmatter_dict.items():
                new_frontmatter += f"{key}: {value}\n"
            new_frontmatter += "---\n\n"

            final_content = content if content is not None else existing_content
            full_content = new_frontmatter + final_content

            doc_path.write_text(full_content, encoding="utf-8")

            doc_info_update = {
                "updated_at": now,
                "size_bytes": len(full_content.encode("utf-8")),
            }

            if title:
                doc_info_update["title"] = title

            if tags is not None:
                doc_info_update["tags"] = tags

            if metadata:
                doc_info_update["metadata"] = {**doc_info.get("metadata", {}), **metadata}

            with self.index_lock:
                self._update_index(doc_id, doc_info_update)

            rel_path = doc_path.relative_to(self.base_path)
            self.git_service.add_files(self.repo_path, [str(rel_path)])
            self.git_service.commit_changes(self.repo_path, commit_message)

            self._update_memory_graph(
                doc_id,
                title or doc_info.get("title"),
                doc_info.get("document_type"),
                tags if tags is not None else doc_info.get("tags", []),
                {**doc_info.get("metadata", {}), **(metadata or {})},
                doc_info.get("source_url"),
            )

            if content is not None:
                self.generate_embeddings(doc_id, content)

        return self.get_document(doc_id)

    def get_document(self, doc_id: str) -> Dict[str, Any]:
        """Get document metadata and preview"""
        doc_info = self._get_document_index(doc_id)
        if not doc_info:
            return None

        doc_path = self.base_path / doc_info["path"]
        if not doc_path.exists():
            return None

        content = doc_path.read_text(encoding="utf-8")

        content_without_frontmatter = re.sub(r"^---.*?---\n\n", "", content, flags=re.DOTALL)
        preview = content_without_frontmatter[:500] + (
            "..." if len(content_without_frontmatter) > 500 else ""
        )

        version_count = 1
        try:
            log = self.git_service.get_log(
                self.repo_path, max_count=100, file_path=doc_info["path"]
            )
            version_count = len(log.get("commits", []))
        except Exception:
            pass

        return {
            "id": doc_id,
            "title": doc_info.get("title", "Untitled"),
            "document_type": doc_info.get("document_type", DocumentType.GENERIC.value),
            "created_at": doc_info.get("created_at", 0),
            "updated_at": doc_info.get("updated_at", 0),
            "tags": doc_info.get("tags", []),
            "metadata": doc_info.get("metadata", {}),
            "content_preview": preview,
            "size_bytes": doc_info.get("size_bytes", 0),
            "version_count": version_count,
            "content_available": True,
            "source_url": doc_info.get("source_url"),
        }

    def get_document_content(self, doc_id: str, version: Optional[str] = None) -> Dict[str, Any]:
        """Get full document content, optionally from a specific version"""
        doc_info = self._get_document_index(doc_id)
        if not doc_info:
            return None

        doc_path = self.base_path / doc_info["path"]

        content = ""
        if version:
            try:
                content = self.git_service.get_file_content_at_version(
                    self.repo_path, doc_info["path"], version
                )
            except Exception:
                return None
        else:
            if not doc_path.exists():
                return None
            content = doc_path.read_text(encoding="utf-8")

        content_without_frontmatter = re.sub(r"^---.*?---\n\n", "", content, flags=re.DOTALL)
        return {
            "id": doc_id,
            "title": doc_info.get("title", "Untitled"),
            "content": content_without_frontmatter,
            "version": version,
        }

    def get_document_versions(self, doc_id: str, max_versions: int = 10) -> List[Dict[str, Any]]:
        """Get version history for a document"""
        doc_info = self._get_document_index(doc_id)
        if not doc_info:
            return []

        try:
            log = self.git_service.get_log(
                self.repo_path, max_count=max_versions, file_path=doc_info["path"]
            )
            versions = []
            for commit in log.get("commits", []):
                versions.append(
                    {
                        "version_hash": commit["hash"],
                        "commit_message": commit["message"],
                        "author": commit["author"],
                        "timestamp": int(
                            time.mktime(time.strptime(commit["date"], "%Y-%m-%d %H:%M:%S %z"))
                        ),
                    }
                )
            return versions
        except Exception as e:
            logger.error(f"Error getting versions: {e}", exc_info=True)
            return []

    def delete_document(self, doc_id: str) -> bool:
        """Delete a document"""
        doc_info = self._get_document_index(doc_id)
        if not doc_info:
            return False

        doc_path = self.base_path / doc_info["path"]
        if not doc_path.exists():
            return False

        try:
            doc_path.unlink()

            rel_path = doc_path.relative_to(self.base_path)
            self.git_service.remove_file(self.repo_path, str(rel_path))
            self.git_service.commit_changes(
                self.repo_path, f"Deleted document: {doc_info.get('title', doc_id)}"
            )

            self._remove_index(doc_id)

            self._remove_from_memory_graph(doc_id)

            return True
        except Exception as e:
            logger.error(f"Error deleting document: {e}", exc_info=True)
            return False

    def search_documents(
        self,
        query: str,
        doc_type: Optional[str] = None,
        tags: Optional[List[str]] = None,
        limit: int = 10,
    ) -> List[Dict[str, Any]]:
        """Search documents by query, type, and tags"""
        results = []

        index_files = list(self.index_dir.glob("*.json"))

        for index_file in index_files:
            try:
                doc_info = json.loads(index_file.read_text(encoding="utf-8"))

                if doc_type and doc_info.get("document_type") != doc_type:
                    continue

                if tags:
                    doc_tags = set(doc_info.get("tags", []))
                    if not all(tag in doc_tags for tag in tags):
                        continue

                if query:
                    query_lower = query.lower()
                    title = doc_info.get("title", "").lower()
                    doc_content = ""

                    if query and not (query_lower in title):
                        try:
                            doc_path = self.base_path / doc_info["path"]
                            content = doc_path.read_text(encoding="utf-8")
                            doc_content = re.sub(
                                r"^---.*?---\n\n", "", content, flags=re.DOTALL
                            ).lower()
                        except Exception:
                            pass

                    if not (query_lower in title or query_lower in doc_content):
                        continue

                results.append(
                    {
                        "id": doc_info.get("id"),
                        "title": doc_info.get("title", "Untitled"),
                        "document_type": doc_info.get("document_type", DocumentType.GENERIC.value),
                        "created_at": doc_info.get("created_at", 0),
                        "updated_at": doc_info.get("updated_at", 0),
                        "tags": doc_info.get("tags", []),
                        "metadata": doc_info.get("metadata", {}),
                        "size_bytes": doc_info.get("size_bytes", 0),
                        "source_url": doc_info.get("source_url"),
                    }
                )

                if len(results) >= limit:
                    break

            except Exception as e:
                logger.error(f"Error processing document index {index_file}: {e}", exc_info=True)

        return results

    def _update_index(self, doc_id: str, doc_info: Dict[str, Any]) -> None:
        """Update the document index"""
        with self.index_lock:
            index_path = self.index_dir / f"{doc_id}.json"

            if index_path.exists():
                current_info = json.loads(index_path.read_text(encoding="utf-8"))
                current_info.update(doc_info)
                index_path.write_text(json.dumps(current_info, indent=2), encoding="utf-8")
            else:
                index_path.write_text(json.dumps(doc_info, indent=2), encoding="utf-8")

    def _get_document_index(self, doc_id: str) -> Optional[Dict[str, Any]]:
        """Get document index by ID"""
        index_path = self.index_dir / f"{doc_id}.json"
        if not index_path.exists():
            return None

        try:
            return json.loads(index_path.read_text(encoding="utf-8"))
        except Exception:
            return None

    def _remove_index(self, doc_id: str) -> None:
        """Remove document index"""
        index_path = self.index_dir / f"{doc_id}.json"
        if index_path.exists():
            index_path.unlink()

    def _update_memory_graph(
        self,
        doc_id: str,
        title: str,
        doc_type: str,
        tags: List[str],
        metadata: Dict[str, Any],
        source_url: Optional[str] = None,
    ) -> None:
        """Update the knowledge graph with document references"""
        try:
            doc_entity_name = f"document:{doc_id}"

            observations = [
                f"Title: {title}",
                f"Type: {doc_type}",
            ]

            if tags:
                observations.append(f"Tags: {', '.join(tags)}")

            if source_url:
                observations.append(f"Source URL: {source_url}")

            for key, value in metadata.items():
                if isinstance(value, (str, int, float, bool)):
                    observations.append(f"{key}: {value}")

            self.memory_service.create_entities(
                [{"name": doc_entity_name, "entity_type": "document", "observations": observations}]
            )

            relations = []
            for tag in tags:
                tag_entity_name = f"tag:{tag}"
                self.memory_service.create_entities(
                    [
                        {
                            "name": tag_entity_name,
                            "entity_type": "tag",
                            "observations": [f"Document tag: {tag}"],
                        }
                    ]
                )

                relations.append(
                    {"from": doc_entity_name, "to": tag_entity_name, "relation_type": "tagged_with"}
                )

            if source_url:
                source_entity_name = f"source:{source_url.replace('://', '_').replace('/', '_')}"
                self.memory_service.create_entities(
                    [
                        {
                            "name": source_entity_name,
                            "entity_type": "source",
                            "observations": [f"URL: {source_url}"],
                        }
                    ]
                )

                relations.append(
                    {
                        "from": doc_entity_name,
                        "to": source_entity_name,
                        "relation_type": "sourced_from",
                    }
                )

            if relations:
                self.memory_service.create_relations(relations)

        except Exception as e:
            logger.error(f"Error updating memory graph: {e}", exc_info=True)

    def _remove_from_memory_graph(self, doc_id: str) -> None:
        """Remove document from knowledge graph"""
        try:
            doc_entity_name = f"document:{doc_id}"
            self.memory_service.delete_entities([doc_entity_name])
        except Exception as e:
            logger.error(f"Error removing document from memory graph: {e}", exc_info=True)

    def generate_embeddings(self, doc_id: str, content: str) -> None:
        """Generate and store embeddings for a document"""
        if not self.vector_search_enabled:
            return

        try:
            embedding = self.vector_model.encode(content[:10000])

            embedding_path = self.vector_index_path / f"{doc_id}.npy"
            self.np.save(embedding_path, embedding)
        except Exception as e:
            logger.error(f"Error generating embeddings for document {doc_id}: {str(e)}")

    def semantic_search(self, query: str, limit: int = 5) -> List[Dict[str, Any]]:
        """Search documents semantically using vector similarity"""
        if not self.vector_search_enabled:
            raise ValueError("Vector search not enabled")

        try:
            query_embedding = self.vector_model.encode(query)

            results = []
            for embedding_file in self.vector_index_path.glob("*.npy"):
                doc_id = embedding_file.stem
                doc_embedding = self.np.load(embedding_file)

                similarity = self.np.dot(query_embedding, doc_embedding) / (
                    self.np.linalg.norm(query_embedding) * self.np.linalg.norm(doc_embedding)
                )

                results.append((doc_id, float(similarity)))

            results.sort(key=lambda x: x[1], reverse=True)
            top_results = results[:limit]

            return [self.get_document(doc_id) for doc_id, _ in top_results]
        except Exception as e:
            logger.error(f"Error during semantic search: {str(e)}")
            return []

    def convert_document_format(self, doc_id: str, target_format: str) -> bytes:
        """Convert document to different formats (PDF, DOCX, etc.)"""
        doc_content = self.get_document_content(doc_id)
        if not doc_content:
            raise ValueError(f"Document {doc_id} not found")

        content = doc_content["content"]
        title = doc_content["title"]

        if target_format.lower() == "pdf":
            try:
                if markdown is None:
                    raise ImportError("markdown package is required for PDF conversion")

                try:
                    import weasyprint
                except ImportError:
                    raise ImportError("weasyprint package is required for PDF conversion")

                html_content = f"<h1>{title}</h1>{markdown.markdown(content)}"
                pdf_bytes = weasyprint.HTML(string=html_content).write_pdf()
                return pdf_bytes
            except ImportError as e:
                logger.error(f"{e}")
                raise ValueError(str(e))

        elif target_format.lower() == "docx":
            try:
                try:
                    from docx import Document
                except ImportError:
                    raise ImportError("python-docx package is required for DOCX conversion")

                doc = Document()
                doc.add_heading(title, 0)
                doc.add_paragraph(content)

                buffer = io.BytesIO()
                doc.save(buffer)
                buffer.seek(0)
                return buffer.read()
            except ImportError as e:
                logger.error(f"{e}")
                raise ValueError(str(e))
        else:
            raise ValueError(f"Unsupported format: {target_format}")

    def get_document_diff(
        self, doc_id: str, from_version: str, to_version: str = "HEAD"
    ) -> Dict[str, Any]:
        """Get differences between document versions"""
        doc_info = self._get_document_index(doc_id)
        if not doc_info:
            raise ValueError(f"Document with ID {doc_id} not found")

        try:
            diff = self.git_service.get_diff(
                self.repo_path, doc_info["path"], from_version, to_version
            )

            return {
                "id": doc_id,
                "title": doc_info.get("title", "Untitled"),
                "from_version": from_version,
                "to_version": to_version,
                "diff": diff,
            }
        except Exception as e:
            logger.error(f"Error getting document diff: {e}", exc_info=True)
            raise ValueError(f"Error getting document diff: {e}")
